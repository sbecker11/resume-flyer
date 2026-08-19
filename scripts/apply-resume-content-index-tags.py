#!/usr/bin/env python3
"""
Prefix job/section headings in a résumé DOCX with [1.1.3] content-index tags.
Idempotent: skips paragraphs that already start with [digits].
"""
from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

CONTENT_INDEX_PREFIX = re.compile(r"^\[\d+(?:\.\d+)*\]\s*")

# Exact heading replacements (employer / section title lines)
HEADING_TAGS: dict[str, str] = {
    "Spexture (Independent Consulting)": "[1] Spexture (Independent Consulting)",
    "Spexture Portfolio Projects": "[1.1] Spexture Portfolio Projects",
    "Spexture Client Engagements": "[1.2] Spexture Client Engagements",
    "SeniorLink (now Careforth)": "[2] SeniorLink (now Careforth)",
    "BigR.io": "[3] BigR.io",
    "ClipFile": "[4] ClipFile",
    "Sierra Vista Group": "[5] Sierra Vista Group",
    "HomePortfolio": "[6] HomePortfolio",
}

# Bullet lines: prefix match (after optional •) → index
BULLET_TAGS: list[tuple[str, str]] = [
    ("resume-parser", "1.1.1"),
    ("color-palette-maker", "1.1.2"),
    ("resume-flyer", "1.1.3"),
    ("sushi-rag-mcp-server", "1.1.4"),
    ("campaign asset automation", "1.1.5"),
    ("spexture-com", "1.1.6"),
    ("ecommerce-semantic-search", "1.1.7"),
    ("linkage-engine", "1.1.8"),
    ("healthcare-agentic-snowflake-rag", "1.1.9"),
    ("recruiting-automation", "1.1.10"),
    ("Adobe (", "1.2.1"),
    ("Fannie Mae (", "1.2.2"),
    ("Cigna (", "1.2.3"),
    ("Warner Bros. Games (", "1.2.4"),
    ("Angel Studios (", "1.2.5"),
    ("Greenseed LLC (", "1.2.6"),
    ("NuSkin (", "1.2.7"),
    ("MSC / One Call (", "5.1"),
    ("Coca-Cola Corp (", "5.2"),
    ("AMI (", "5.3"),
    ("Intrusic (", "5.4"),
    ("Rowe International (", "5.6"),
]

RECRUITING_LINE = (
    "[1.1.10] • recruiting-automation — launchd-scheduled pipeline orchestrating "
    "comms-migration inbox classification and job-tracker Gmail triage, LLM JD evaluation, "
    "and tailored résumé/cover-letter package generation "
    "(github.com/sbecker11/recruiting-automation)"
)


def _strip_bullet(text: str) -> str:
    return re.sub(r"^[\u2022\u00b7•]\s*", "", text.strip())


def tag_paragraph_text(text: str) -> str | None:
    stripped = text.strip()
    if not stripped or CONTENT_INDEX_PREFIX.match(stripped):
        return None

    if stripped in HEADING_TAGS:
        return HEADING_TAGS[stripped]

    body = _strip_bullet(stripped)
    for needle, index in BULLET_TAGS:
        if body.startswith(needle) or needle in body[: len(needle) + 2]:
            if stripped.startswith("•") or stripped.startswith("·"):
                return f"[{index}] {stripped}"
            return f"[{index}] • {stripped}"
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para


def apply_tags(doc_path: Path, out_path: Path | None = None) -> int:
    doc = Document(doc_path)
    changed = 0
    insert_after_idx: int | None = None

    for i, para in enumerate(doc.paragraphs):
        new_text = tag_paragraph_text(para.text)
        if new_text:
            para.text = new_text
            changed += 1
        if "healthcare-agentic-snowflake-rag" in para.text and "1.1.9" in para.text:
            insert_after_idx = i

    has_recruiting = any("recruiting-automation" in p.text for p in doc.paragraphs)
    if not has_recruiting and insert_after_idx is not None:
        insert_paragraph_after(doc.paragraphs[insert_after_idx], RECRUITING_LINE)
        changed += 1

    target = out_path or doc_path
    doc.save(target)
    return changed


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: apply-resume-content-index-tags.py <input.docx> [output.docx]", file=sys.stderr)
        return 1
    src = Path(sys.argv[1]).expanduser()
    dst = Path(sys.argv[2]).expanduser() if len(sys.argv) > 2 else src
    if not src.is_file():
        print(f"Error: not found: {src}", file=sys.stderr)
        return 1
    n = apply_tags(src, dst if dst != src else None)
    print(f"Tagged {n} paragraph(s); wrote {dst if dst != src else src}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
